# -*- coding: utf-8 -*-
"""
Dashboard Operacional — Vistorias CRO/1
Versão adaptada para estrutura modular do projeto IME / CRO1
Autor: Robson Nunes Rodrigues Junior
Data: Atualizado em 13/10/2025
"""

from __future__ import annotations
import io
from datetime import datetime
import pandas as pd
import numpy as np
import plotly.express as px
import streamlit as st

# Importações do core do projeto
from core.data_loader import read_df
from core.config import TAB_SOLICITACOES

# =========================================================
# Funções auxiliares
# =========================================================
def _nf(s: str) -> str:
    return (s or "").casefold().strip()

def _pick(df: pd.DataFrame, candidates: list[str]) -> str | None:
    """Busca coluna por nome aproximado (case insensitive)."""
    if df is None or df.empty:
        return None
    for c in candidates:
        for col in df.columns:
            if _nf(col) == _nf(c) or _nf(c) in _nf(col):
                return col
    return None

def _to_date(x):
    if pd.isna(x):
        return pd.NaT
    if isinstance(x, (datetime, pd.Timestamp)):
        return pd.to_datetime(x)
    try:
        return pd.to_datetime(x, dayfirst=True, errors="coerce")
    except Exception:
        return pd.NaT

def _safe_num(x):
    try:
        if pd.isna(x):
            return np.nan
        if isinstance(x, str):
            return float(x.replace("R$", "").replace(".", "").replace(",", ".").strip())
        return float(x)
    except Exception:
        return np.nan

# =========================================================
# Função principal do Dashboard
# =========================================================
def main():
    st.set_page_config(
        page_title="Dashboard CRO/1 — Vistorias",
        layout="wide",
        page_icon="📊"
    )

    st.title("📊 Dashboard Operacional — CRO/1 (Vistorias Técnicas)")

    # -----------------------------------------------------
    # Leitura da base
    # -----------------------------------------------------
    try:
        df_raw = read_df(TAB_SOLICITACOES)
    except Exception as e:
        st.error(f"Erro ao carregar a base de dados: {e}")
        return

    if df_raw is None or df_raw.empty:
        st.warning("Não há registros na base de solicitações (TAB_SOLICITACOES).")
        return

    df = df_raw.copy()

    # -----------------------------------------------------
    # Identificação automática de colunas
    # -----------------------------------------------------
    COL = {
        "om": _pick(df, ["OM", "OM beneficiada", "OM apoiada", "Organização Militar"]),
        "diretoria": _pick(df, ["Diretoria", "Dir responsável", "DIR", "Direção"]),
        "especialidade": _pick(df, [
            "Especialidade", "Especialidade envolvida", "Tipo/Especialidade",
            "Engenharia", "Área técnica", "Filtro - qual a especialidade da VT"
        ]),
        "prioridade": _pick(df, [
            "Prioridade", "Classificação", "Tratativa da vistoria", "Classe da demanda",
            "Normal/Prioridade/Urgente/Urgentíssimo"
        ]),
        "status": _pick(df, ["Status da Vistoria", "Status", "Situação", "Andamento"]),
        "dt_solic": _pick(df, ["Data da solicitação", "Dt Solicitação", "Solicitado em"]),
        "dt_real_visita": _pick(df, ["Data da realização da vistoria", "Data da vistoria", "Realização da visita", "Data visita"]),
        "dt_conc": _pick(df, ["Data da conclusão da VT", "Data da conclusão", "Conclusão da VT", "Conclusão"]),
        "orcamento": _pick(df, ["Orçamento estimado", "Valor estimado", "Custo", "PFR", "Total R$", "Orçamento"]),
    }
    
    # normalizações
    for k in ["dt_solic", "dt_real_visita", "dt_conc"]:
        if COL[k]:
            df[COL[k]] = df[COL[k]].map(_to_date)
    
    if COL["orcamento"]:
        df[COL["orcamento"]] = df[COL["orcamento"]].map(_safe_num)
    
    # Campo de emergência (seguro mesmo sem coluna)
    if COL["prioridade"] and COL["prioridade"] in df.columns:
        df["Classificação"] = np.where(
            df[COL["prioridade"]].astype(str).str.contains("urg|emerg", case=False, na=False),
            "Emergencial", "Não Emergencial"
        )
    else:
        df["Classificação"] = "Não Informado"

    # =========================================================
    # Filtros laterais
    # =========================================================
    with st.sidebar:
        st.header("Filtros")
    
        # Período
        if COL["dt_solic"] and COL["dt_solic"] in df.columns:
            min_d = pd.to_datetime(df[COL["dt_solic"]]).min()
            max_d = pd.to_datetime(df[COL["dt_solic"]]).max()
            periodo = st.date_input(
                "Período",
                (min_d.date() if pd.notna(min_d) else datetime(2025,1,1).date(),
                 max_d.date() if pd.notna(max_d) else datetime.today().date())
            )
        else:
            periodo = (datetime(2025,1,1).date(), datetime.today().date())
    
        def _opts(key):
            col = COL.get(key)
            if col and col in df.columns:
                return sorted(df[col].dropna().astype(str).unique())
            return []
    
        om_sel  = st.multiselect("OM", _opts("om"))
        esp_sel = st.multiselect("Especialidade", _opts("especialidade"))
        stat_sel= st.multiselect("Status", _opts("status"))
    
    # aplica filtros
    mask = pd.Series(True, index=df.index)
    
    if COL["dt_solic"] and COL["dt_solic"] in df.columns:
        mask &= df[COL["dt_solic"]].between(pd.to_datetime(periodo[0]), pd.to_datetime(periodo[1]))
    
    if om_sel and COL["om"] and COL["om"] in df.columns:
        mask &= df[COL["om"]].astype(str).isin(om_sel)
    
    if esp_sel and COL["especialidade"] and COL["especialidade"] in df.columns:
        mask &= df[COL["especialidade"]].astype(str).isin(esp_sel)
    
    if stat_sel and COL["status"] and COL["status"] in df.columns:
        mask &= df[COL["status"]].astype(str).isin(stat_sel)
    
    dff = df[mask].copy()

    # =========================================================
    # Derivados (prazo) — criados ANTES dos KPIs e do DOCX
    # =========================================================
    if (COL["dt_solic"] and COL["dt_solic"] in dff.columns) and (COL["dt_conc"] and COL["dt_conc"] in dff.columns):
        dff["Tempo Execução (dias)"] = (dff[COL["dt_conc"]] - dff[COL["dt_solic"]]).dt.days

    # =========================================================
    # Indicadores principais (KPIs)
    # =========================================================
    st.subheader("Indicadores de Desempenho")
    col1, col2, col3, col4 = st.columns(4)
    
    col1.metric("Total de Vistorias", len(dff))
    
    # % emergenciais (só conta se houver "Emergencial"/"Não Emergencial")
    if "Classificação" in dff.columns:
        base_emerg = dff["Classificação"].isin(["Emergencial","Não Emergencial"])
        pct_emerg = 100 * dff.loc[base_emerg, "Classificação"].eq("Emergencial").mean() if base_emerg.any() else np.nan
        col2.metric("% Emergenciais", f"{pct_emerg:.1f}%" if pd.notna(pct_emerg) else "—")
    else:
        col2.metric("% Emergenciais", "—")
    
    if COL["orcamento"] and COL["orcamento"] in dff.columns:
        total_rs = dff[COL["orcamento"]].sum()
        col3.metric("Orçamento Total", f"R$ {total_rs:,.2f}".replace(",", "X").replace(".", ",").replace("X","."))
    else:
        col3.metric("Orçamento Total", "—")
    
    if "Tempo Execução (dias)" in dff.columns and dff["Tempo Execução (dias)"].notna().any():
        col4.metric("Prazo Médio de Execução", f"{dff['Tempo Execução (dias)'].mean():.1f} dias")
    else:
        col4.metric("Prazo Médio de Execução", "—")

    # =========================================================
    # Gráficos principais
    # =========================================================
    st.subheader("Visualizações Analíticas")

    # 1. Série temporal
    if COL["dt_solic"] and COL["dt_solic"] in dff.columns:
        df_mes = (dff.groupby(dff[COL["dt_solic"]].dt.to_period("M").dt.to_timestamp())
                     .size().reset_index(name="qtd"))
        fig1 = px.line(df_mes, x=COL["dt_solic"], y="qtd", title="Evolução Mensal das Vistorias", markers=True)
        fig1.update_traces(line_shape="spline")
        st.plotly_chart(fig1, use_container_width=True)
    
    # 2. Status
    if COL["status"] and COL["status"] in dff.columns:
        fig2 = px.pie(dff, names=COL["status"], title="Distribuição por Status", hole=0.45)
        st.plotly_chart(fig2, use_container_width=True)
    
    # 3. Top OMs
    if COL["om"] and COL["om"] in dff.columns:
        top_om = dff.groupby(COL["om"]).size().nlargest(10).reset_index(name="Qtd")
        fig3 = px.bar(top_om, x="Qtd", y=COL["om"], orientation="h", title="Top 10 OMs — Quantidade de Vistorias")
        st.plotly_chart(fig3, use_container_width=True)
    
    # 4. Orçamento por classificação
    if (COL["orcamento"] and COL["orcamento"] in dff.columns) and ("Classificação" in dff.columns):
        by_class = dff.groupby("Classificação")[COL["orcamento"]].sum().reset_index()
        fig4 = px.bar(by_class, x="Classificação", y=COL["orcamento"], text_auto=".2s", title="Orçamento por Classificação")
        st.plotly_chart(fig4, use_container_width=True)

    # Aviso de colunas ausentes
    missing = [k for k,v in COL.items() if k in ["om","status","especialidade","dt_solic"] and not v]
    if missing:
        st.warning("Colunas não encontradas na base: " + ", ".join(missing) +
                   ". O dashboard continua funcionando, mas alguns filtros/gráficos serão ocultados.")

    # =========================================================
    # Tabela detalhada
    # =========================================================
    st.subheader("Tabela Detalhada de Vistorias")
    st.dataframe(dff, use_container_width=True, hide_index=True)

    # =========================================================
    # Exportação
    # =========================================================
    st.subheader("Exportar Relatório Resumido")

    if st.button("Gerar DOCX"):
        from docx import Document
        doc = Document()
        doc.add_heading("Relatório Resumido — Vistorias CRO/1", 0)
        doc.add_paragraph(f"Gerado em {datetime.now():%d/%m/%Y %H:%M}")

        doc.add_heading("Indicadores Principais", level=2)
        doc.add_paragraph(f"Total de Vistorias: {len(dff)}")
        doc.add_paragraph(f"% Emergenciais: {100 * dff['Classificação'].eq('Emergencial').mean():.1f}%")
        if COL["orcamento"] and COL["orcamento"] in dff.columns:
            doc.add_paragraph(
                f"Orçamento Total: R$ {dff[COL['orcamento']].sum():,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            )
        if "Tempo Execução (dias)" in dff.columns:
            doc.add_paragraph(f"Prazo Médio de Execução: {dff['Tempo Execução (dias)'].mean():.1f} dias")

        doc.add_heading("Tabela de Vistorias", level=2)
        t = doc.add_table(rows=1, cols=len(dff.columns))
        hdr_cells = t.rows[0].cells
        for j, c in enumerate(dff.columns):
            hdr_cells[j].text = c
        for _, row in dff.head(25).iterrows():
            row_cells = t.add_row().cells
            for j, c in enumerate(dff.columns):
                row_cells[j].text = str(row[c]) if pd.notna(row[c]) else ""

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button(
            "📄 Baixar Relatório DOCX",
            data=buffer,
            file_name="Relatorio_Vistorias_CRO1.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# =========================================================
# Execução direta
# =========================================================
if __name__ == "__main__":
    main()
