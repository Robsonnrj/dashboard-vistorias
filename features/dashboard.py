# -*- coding: utf-8 -*-
"""
features/dashboard.py
Dashboard Operacional — Vistorias CRO/1
"""

from __future__ import annotations
from datetime import datetime
import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st

from core.data_loader import read_df
from core.config import TAB_SOLICITACOES
from core.utils import pick_col


# ----------------------- Helpers -----------------------
def _mes_label(s):
    s = pd.to_datetime(s, errors="coerce")
    return s.dt.to_period("M").dt.to_timestamp()

def _fmt_rs(v):
    if pd.isna(v): 
        return "—"
    return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def _to_date(x):
    """Converte para Timestamp (brasil: dayfirst=True), devolvendo NaT em erros."""
    if pd.isna(x):
        return pd.NaT
    if isinstance(x, (datetime, pd.Timestamp)):
        return pd.to_datetime(x, errors="coerce")
    return pd.to_datetime(x, errors="coerce", dayfirst=True)

def _safe_num(x):
    """Converte valores monetários brasileiros para float."""
    try:
        if pd.isna(x):
            return np.nan
        if isinstance(x, str):
            return float(x.replace("R$", "").replace(".", "").replace(",", ".").strip())
        return float(x)
    except Exception:
        return np.nan

def _safestr(x) -> str:
    """Converte qualquer valor para str sem estourar com pd.NA/NaT."""
    return "" if pd.isna(x) else str(x)


# ----------------------- Main page Feature -----------------------
def page():
    st.caption("Use os filtros ao lado para refinar os indicadores.")
    st.header("📊 Dashboard Operacional — Seção de Vistorias")

    # ---- Carregamento principal ----
    try:
        df = read_df(TAB_SOLICITACOES)
    except Exception as e:
        st.error(f"Erro ao carregar a base de dados: {e}")
        return

    if df is None or df.empty:
        st.warning("Não há registros na base de solicitações.")
        return

    # ---- Mapeamento tolerante das colunas principais ----
    COL = {
        "om": pick_col(df, ["OM", "OM beneficiada", "OM apoiada", "Organização Militar", "OM APOIADA"]),
        "diretoria": pick_col(df, ["Diretoria", "Dir responsável", "DIR", "Direção", "Diretoria Responsável"]),
        "especialidade": pick_col(df, [
            "Especialidade", "Especialidade envolvida", "Tipo/Especialidade",
            "Engenharia", "Área técnica", "Filtro - qual a especialidade da VT"
        ]),
        "prioridade": pick_col(df, [
            "Prioridade", "Classificação", "Tratativa da vistoria", "Classe da demanda",
            "Normal/Prioridade/Urgente/Urgentíssimo"
        ]),
        "status": pick_col(df, ["Status da Vistoria", "Status", "Situação", "Andamento"]),
        "dt_solic": pick_col(df, ["Data da solicitação", "Dt Solicitação", "Solicitado em", "DATA DA SOLICITAÇÃO"]),
        "dt_real_visita": pick_col(df, [
            "Data da realização da vistoria", "Data da vistoria", "Realização da visita",
            "Data visita", "DATA DA VISTORIA"
        ]),
        "dt_conc": pick_col(df, [
            "Data da conclusão da VT", "Data da conclusão", "Conclusão da VT", "Conclusão",
            "DATA/PREVISÃO DE CONCLUSÃO"
        ]),
        "dt_resp": pick_col(df, ["DATA DA RESPOSTA A SOLICITAÇÃO", "Data da resposta"]),
        "orcamento": pick_col(df, ["Orçamento estimado", "Valor estimado", "Custo", "PFR", "Total R$", "Orçamento"]),
    }

    # ---- Normalização de datas e valores ----
    for k in ["dt_solic", "dt_real_visita", "dt_conc", "dt_resp"]:
        if COL[k]:
            df[COL[k]] = df[COL[k]].map(_to_date)

    if COL["orcamento"]:
        df[COL["orcamento"]] = df[COL["orcamento"]].map(_safe_num)

    # ---- Emergencialidade (KPI antigo) ----
    if COL["prioridade"]:
        df["Classificação"] = np.where(
            df[COL["prioridade"]].astype(str).str.contains("urg|emerg", case=False, na=False),
            "Emergencial", "Não Emergencial"
        )
    else:
        df["Classificação"] = "Não Informado"

    # ----------------------- Filtros (robustos) -----------------------
    with st.sidebar:
        st.header("Filtros")

        # Reset rápido
        if st.button("🔄 Limpar filtros"):
            for k in ("om", "especialidade", "status", "periodo"):
                st.session_state.pop(k, None)

        incluir_sem_data = False
        if COL["dt_solic"]:
            serie_datas = pd.to_datetime(df[COL["dt_solic"]], errors="coerce", dayfirst=True)
            min_d = serie_datas.min()
            max_d = serie_datas.max()
            if pd.isna(min_d) or pd.isna(max_d):
                min_d = pd.Timestamp(2025, 1, 1)
                max_d = pd.Timestamp.today()

            periodo = st.date_input(
                "Período (Data da Solicitação)",
                key="periodo",
                value=(min_d.date(), max_d.date()),
            )
            incluir_sem_data = st.checkbox("Incluir vistorias sem data de solicitação", value=True)
        else:
            periodo = (datetime(2025, 1, 1).date(), datetime.today().date())

        def _opts(key):
            col = COL.get(key)
            if col and col in df.columns:
                return sorted(df[col].dropna().astype(str).unique())
            return []

        om_sel  = st.multiselect("OM", _opts("om"), key="om")
        esp_sel = st.multiselect("Especialidade", _opts("especialidade"), key="especialidade")
        stat_sel= st.multiselect("Status", _opts("status"), key="status")

    # Aplicação dos filtros
    mask = pd.Series(True, index=df.index, dtype=bool)

    if COL["dt_solic"] and COL["dt_solic"] in df.columns:
        datas = pd.to_datetime(df[COL["dt_solic"]], errors="coerce", dayfirst=True)
        no_intervalo = datas.between(
            pd.to_datetime(periodo[0]), pd.to_datetime(periodo[1]), inclusive="both"
        )
        sem_data = datas.isna()
        if incluir_sem_data:
            mask &= (no_intervalo | sem_data)
        else:
            mask &= no_intervalo.fillna(False)

    if om_sel and COL["om"]:
        mask &= df[COL["om"]].astype(str).isin(om_sel).fillna(False)
    if esp_sel and COL["especialidade"]:
        mask &= df[COL["especialidade"]].astype(str).isin(esp_sel).fillna(False)
    if stat_sel and COL["status"]:
        mask &= df[COL["status"]].astype(str).isin(stat_sel).fillna(False)

    dff = df[mask].copy()

    # Prazo de execução (antigo)
    if COL["dt_solic"] and COL["dt_conc"]:
        dff["Tempo Execução (dias)"] = (dff[COL["dt_conc"]] - dff[COL["dt_solic"]]).dt.days

    # ----------------------- Diagnóstico dos filtros -----------------------
    with st.expander("🔎 Diagnóstico dos filtros (clique para abrir)"):
        total_base = len(df)
        total_filtrado = len(dff)
        st.write(f"Total na base: **{total_base}** | Após filtros: **{total_filtrado}**")
        if COL["dt_solic"]:
            datas_all = pd.to_datetime(df[COL["dt_solic"]], errors="coerce", dayfirst=True)
            fora_periodo = datas_all[
                ~datas_all.between(pd.to_datetime(periodo[0]), pd.to_datetime(periodo[1]), inclusive="both")
                & datas_all.notna()
            ]
            qtd_nat = datas_all.isna().sum()
            st.write(f"- Linhas **sem data de solicitação (NaT)**: **{qtd_nat}**")
            st.write(f"- Linhas **com data fora do período**: **{fora_periodo.shape[0]}**")
            if qtd_nat > 0 and not incluir_sem_data:
                st.info("Há vistorias **sem data**. Marque 'Incluir vistorias sem data de solicitação' para exibi-las.")

    # ============================================================
    # ========== PARÂMETROS DO BILHETE (PEDIDOS/ATENDIDAS) =======
    # ============================================================

    # regras: o que é atendida?
    def _eh_atendida(row) -> bool:
        textos = []
        if COL["status"]:
            textos.append(_safestr(row.get(COL["status"])))
        if "STATUS - ATUALIZAÇÃO SEMANAL" in dff.columns:
            textos.append(_safestr(row.get("STATUS - ATUALIZAÇÃO SEMANAL")))
        j = " ".join(textos).lower()

        if any(p in j for p in ("finaliz", "conclu", "atendid")):
            return True
        if COL["dt_conc"] and pd.notna(row.get(COL["dt_conc"], pd.NaT)):
            return True
        if COL["dt_resp"] and pd.notna(row.get(COL["dt_resp"], pd.NaT)):
            return True
        if COL["dt_real_visita"] and pd.notna(row.get(COL["dt_real_visita"], pd.NaT)):
            return True
        return False

    # flags
    if COL["dt_solic"]:
        dff["_pedido"] = dff[COL["dt_solic"]].notna()
        dff["_mes"] = dff[COL["dt_solic"]].dt.to_period("M").dt.to_timestamp()
    else:
        dff["_pedido"] = True
        dff["_mes"] = pd.NaT

    dff["_atendida"] = dff.apply(_eh_atendida, axis=1)
    dff["_nao_atendida"] = dff["_pedido"] & (~dff["_atendida"])

    # data fim para tempo de atendimento
    def _data_fim(row):
        for k in ("dt_conc", "dt_resp", "dt_real_visita"):
            c = COL.get(k)
            if c and pd.notna(row.get(c, pd.NaT)):
                return row[c]
        return pd.NaT

    dff["_fim"] = dff.apply(_data_fim, axis=1)
    if COL["dt_solic"]:
        dff["_tempo_dias"] = (dff["_fim"] - dff[COL["dt_solic"]]).dt.days
    else:
        dff["_tempo_dias"] = np.nan

    # KPIs do bilhete
    st.subheader("Indicadores de Desempenho (atualizados)")
    c1, c2, c3, c4, c5 = st.columns(5)

    tot_ped = int(dff["_pedido"].sum())
    tot_atd = int(dff["_atendida"].sum())
    pct_atd = (100 * tot_atd / tot_ped) if tot_ped else np.nan
    tempo_medio = dff.loc[dff["_tempo_dias"].notna(), "_tempo_dias"].mean()

    # série mensal
    mensal = (
        dff.loc[dff["_mes"].notna()]
           .groupby("_mes")
           .agg(pedidos=("_pedido", "sum"), atendidas=("_atendida", "sum"))
           .reset_index()
           .sort_values("_mes")
    )
    mensal["backlog"] = (mensal["pedidos"] - mensal["atendidas"]).cumsum()

    c1.metric("Pedidos (solicitações)", f"{tot_ped}")
    c2.metric("Atendidas", f"{tot_atd}")
    c3.metric("% Atendimento", f"{pct_atd:.1f}%" if pd.notna(pct_atd) else "—")
    c4.metric("Tempo médio de atendimento", f"{tempo_medio:.1f} dias" if pd.notna(tempo_medio) else "—")
    c5.metric("Backlog acumulado", int(mensal["backlog"].iloc[-1]) if not mensal.empty else 0)

    # Linha: Pedidos x Atendidas (+ backlog)
    st.subheader("Evolução mensal — Pedidos x Atendidas")
    if COL["dt_solic"] and not mensal.empty:
        fig = go.Figure()
        fig.add_trace(go.Scatter(x=mensal["_mes"], y=mensal["pedidos"],
                                 mode="lines+markers", name="Pedidos"))
        fig.add_trace(go.Scatter(x=mensal["_mes"], y=mensal["atendidas"],
                                 mode="lines+markers", name="Atendidas"))
        fig.add_trace(go.Scatter(x=mensal["_mes"], y=mensal["backlog"],
                                 mode="lines", name="Backlog (acum.)"))
        fig.update_layout(xaxis_title="Mês", yaxis_title="Qtd")
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("Sem dados de 'Data da solicitação' suficientes para a série temporal.")

    # Colunas empilhadas: Diretoria — Atendidas x Não atendidas
    st.subheader("Diretoria — Atendidas x Não atendidas")
    if COL["diretoria"]:
        base_dir = (
            dff.groupby(COL["diretoria"])
               .agg(Atendidas=("_atendida", "sum"), Nao_Atendidas=("_nao_atendida", "sum"))
               .reset_index()
        )
        base_dir = base_dir.sort_values(["Atendidas", "Nao_Atendidas"], ascending=[False, False])
        fig = go.Figure()
        fig.add_bar(x=base_dir[COL["diretoria"]], y=base_dir["Atendidas"], name="Atendidas")
        fig.add_bar(x=base_dir[COL["diretoria"]], y=base_dir["Nao_Atendidas"], name="Não atendidas")
        fig.update_layout(barmode="stack", xaxis_title="Diretoria", yaxis_title="Qtd")
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("Coluna de Diretoria não encontrada.")

    # Pizza: Atendidas x Não atendidas
    st.subheader("Participação — Atendidas x Não atendidas")
    pie = pd.DataFrame({
        "Categoria": ["Atendidas", "Não atendidas"],
        "Qtd": [int(dff["_atendida"].sum()), int(dff["_nao_atendida"].sum())]
    })
    fig = px.pie(pie, names="Categoria", values="Qtd", hole=0.45)
    st.plotly_chart(fig, use_container_width=True)

    # Radar — tempo médio por OM (top 10)
    st.subheader("Radar — Tempo médio de atendimento por OM (top 10)")
    if COL["om"]:
        rad = (
            dff.loc[dff["_tempo_dias"].notna()]
               .groupby(COL["om"])["_tempo_dias"].mean()
               .sort_values(ascending=False).head(10).reset_index()
        )
        if not rad.empty:
            fig = go.Figure()
            fig.add_trace(go.Scatterpolar(r=rad["_tempo_dias"], theta=rad[COL["om"]],
                                          fill="toself", name="dias"))
            fig.update_layout(polar=dict(radialaxis=dict(visible=True)), showlegend=False)
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("Ainda não há pares data_solicitação x data_fim suficientes para o radar.")
    else:
        st.info("Coluna de OM não encontrada.")

    # Mantido: Diretoria por OM — quantidade
    st.subheader("Diretoria por OM — quantidade de vistorias")
    if COL["om"] and COL["diretoria"]:
        top = (dff.groupby([COL["diretoria"], COL["om"]]).size()
                 .reset_index(name="Qtd")
                 .sort_values("Qtd", ascending=False)
                 .head(15))
        fig = px.bar(top, x="Qtd", y=COL["om"], color=COL["diretoria"], orientation="h",
                     title="Top OMs por Diretoria (Qtd de vistorias)")
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("Preciso das colunas Diretoria e OM.")

    # ----------------------- Visualizações (as suas originais) -----------------------
    st.subheader("Visualizações Analíticas")

    if COL["dt_solic"]:
        dff[COL["dt_solic"]] = pd.to_datetime(dff[COL["dt_solic"]], errors="coerce", dayfirst=True)
        df_mes = (
            dff.groupby(dff[COL["dt_solic"]].dt.to_period("M").dt.to_timestamp())
               .size()
               .reset_index(name="qtd")
               .sort_values(COL["dt_solic"])
        )
        fig1 = px.line(df_mes, x=COL["dt_solic"], y="qtd",
                       title="Evolução Mensal das Vistorias", markers=True)
        fig1.update_traces(line_shape="spline")
        st.plotly_chart(fig1, use_container_width=True)

    if COL["status"]:
        fig2 = px.pie(dff, names=COL["status"], title="Distribuição por Status", hole=0.45)
        st.plotly_chart(fig2, use_container_width=True)

    if COL["om"]:
        top_om = dff.groupby(COL["om"]).size().nlargest(10).reset_index(name="Qtd")
        fig3 = px.bar(top_om, x="Qtd", y=COL["om"], orientation="h",
                      title="Top 10 OMs — Quantidade de Vistorias")
        st.plotly_chart(fig3, use_container_width=True)

    if COL["orcamento"] and "Classificação" in dff.columns:
        by_class = dff.groupby("Classificação")[COL["orcamento"]].sum().reset_index()
        fig4 = px.bar(by_class, x="Classificação", y=COL["orcamento"], text_auto=".2s",
                      title="Orçamento por Classificação")
        st.plotly_chart(fig4, use_container_width=True)

    # Aviso de colunas faltantes
    missing = [k for k, v in COL.items() if k in ["om", "status", "especialidade", "dt_solic"] and not v]
    if missing:
        st.warning(
            "Colunas não encontradas na base: " + ", ".join(missing)
            + ". O dashboard continua funcionando, mas alguns filtros/gráficos serão ocultados."
        )

    # ----------------------- Tabela e Exportação -----------------------
    st.subheader("Tabela Detalhada de Vistorias")
    st.dataframe(dff, use_container_width=True, hide_index=True)

    st.subheader("Exportar Relatório Resumido")
    if st.button("Gerar DOCX"):
        try:
            from docx import Document
        except ModuleNotFoundError:
            st.error(
                "Pacote **python-docx** não está instalado. "
                "Adicione `python-docx==0.8.11` ao seu `requirements.txt` e faça o deploy novamente."
            )
        else:
            doc = Document()
            doc.add_heading("Relatório Resumido — Vistorias CRO/1", 0)
            doc.add_paragraph(f"Gerado em {datetime.now():%d/%m/%Y %H:%M}")

            doc.add_heading("Indicadores Principais", level=2)
            doc.add_paragraph(f"Total de Vistorias: {len(dff)}")
            if "Classificação" in dff.columns and not dff.empty:
                doc.add_paragraph(f"% Emergenciais: {100 * dff['Classificação'].eq('Emergencial').mean():.1f}%")
            if COL["orcamento"] and COL["orcamento"] in dff.columns:
                doc.add_paragraph(
                    f"Orçamento Total: R$ {dff[COL['orcamento']].sum():,.2f}"
                    .replace(",", "X").replace(".", ",").replace("X", ".")
                )
            if "Tempo Execução (dias)" in dff.columns:
                doc.add_paragraph(f"Prazo Médio de Execução: {dff['Tempo Execução (dias)'].mean():.1f} dias")

            import io
            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
            st.download_button(
                "📄 Baixar Relatório DOCX",
                data=bio,
                file_name="Relatorio_Vistorias_CRO1.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
